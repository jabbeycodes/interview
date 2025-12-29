"""
Resume parser - extracts structured data from PDF and DOCX resumes
"""
import io
import re
from typing import Optional, Dict, List, Tuple
from pathlib import Path
from datetime import datetime


class ResumeParser:
    """
    Parses resume files (PDF, DOCX) and extracts structured information.
    """

    # Common section headers
    SECTION_HEADERS = {
        "experience": [
            "experience", "work experience", "professional experience",
            "employment history", "work history", "career history"
        ],
        "education": [
            "education", "academic background", "educational background",
            "qualifications", "academic qualifications"
        ],
        "skills": [
            "skills", "technical skills", "core competencies",
            "expertise", "technologies", "proficiencies"
        ],
        "projects": [
            "projects", "personal projects", "key projects",
            "selected projects", "portfolio"
        ],
        "summary": [
            "summary", "professional summary", "profile",
            "about me", "objective", "career objective"
        ],
        "certifications": [
            "certifications", "certificates", "licenses",
            "professional certifications"
        ]
    }

    def __init__(self):
        self._pdf_available = self._check_pdf_support()
        self._docx_available = self._check_docx_support()

    def _check_pdf_support(self) -> bool:
        """Check if PDF parsing is available"""
        try:
            import PyPDF2
            return True
        except ImportError:
            try:
                import pdfplumber
                return True
            except ImportError:
                return False

    def _check_docx_support(self) -> bool:
        """Check if DOCX parsing is available"""
        try:
            import docx
            return True
        except ImportError:
            return False

    def parse_file(self, file_path: str) -> Dict:
        """
        Parse a resume file and extract structured data.

        Args:
            file_path: Path to the resume file (PDF or DOCX)

        Returns:
            Dictionary with extracted resume sections
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            text = self._extract_pdf_text(file_path)
        elif suffix in [".docx", ".doc"]:
            text = self._extract_docx_text(file_path)
        elif suffix == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

        return self._parse_text(text)

    def parse_bytes(self, file_bytes: bytes, filename: str) -> Dict:
        """
        Parse resume from bytes.

        Args:
            file_bytes: Raw file bytes
            filename: Original filename for format detection

        Returns:
            Dictionary with extracted resume sections
        """
        suffix = Path(filename).suffix.lower()

        if suffix == ".pdf":
            text = self._extract_pdf_text_from_bytes(file_bytes)
        elif suffix in [".docx", ".doc"]:
            text = self._extract_docx_text_from_bytes(file_bytes)
        elif suffix == ".txt":
            text = file_bytes.decode("utf-8")
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

        return self._parse_text(text)

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() or ""
                    text += "\n"
                return text
        except ImportError:
            pass

        try:
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                    text += "\n"
                return text
        except ImportError:
            raise ImportError("Please install pdfplumber or PyPDF2: pip install pdfplumber")

    def _extract_pdf_text_from_bytes(self, file_bytes: bytes) -> str:
        """Extract text from PDF bytes"""
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() or ""
                    text += "\n"
                return text
        except ImportError:
            pass

        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
                text += "\n"
            return text
        except ImportError:
            raise ImportError("Please install pdfplumber or PyPDF2")

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            import docx
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            raise ImportError("Please install python-docx: pip install python-docx")

    def _extract_docx_text_from_bytes(self, file_bytes: bytes) -> str:
        """Extract text from DOCX bytes"""
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            raise ImportError("Please install python-docx: pip install python-docx")

    def _parse_text(self, text: str) -> Dict:
        """
        Parse extracted text into structured sections.

        Args:
            text: Raw text from resume

        Returns:
            Dictionary with structured resume data
        """
        result = {
            "raw_text": text,
            "name": None,
            "email": None,
            "phone": None,
            "linkedin": None,
            "summary": None,
            "experiences": [],
            "education": [],
            "skills": [],
            "projects": [],
            "certifications": []
        }

        # Clean text
        text = self._clean_text(text)
        lines = text.split("\n")

        # Extract contact info
        result["email"] = self._extract_email(text)
        result["phone"] = self._extract_phone(text)
        result["linkedin"] = self._extract_linkedin(text)
        result["name"] = self._extract_name(lines)

        # Identify sections
        sections = self._identify_sections(lines)

        # Parse each section
        if "summary" in sections:
            result["summary"] = self._parse_summary(sections["summary"])

        if "experience" in sections:
            result["experiences"] = self._parse_experiences(sections["experience"])

        if "education" in sections:
            result["education"] = self._parse_education(sections["education"])

        if "skills" in sections:
            result["skills"] = self._parse_skills(sections["skills"])

        if "projects" in sections:
            result["projects"] = self._parse_projects(sections["projects"])

        return result

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Restore line breaks
        text = re.sub(r'([.!?])\s+([A-Z])', r'\1\n\2', text)
        # Clean up
        text = text.strip()
        return text

    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address"""
        pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
        match = re.search(pattern, text)
        return match.group(0) if match else None

    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number"""
        patterns = [
            r'\+?1?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\d{3}[-.\s]\d{3}[-.\s]\d{4}'
        ]
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        return None

    def _extract_linkedin(self, text: str) -> Optional[str]:
        """Extract LinkedIn URL"""
        pattern = r'linkedin\.com/in/[\w-]+'
        match = re.search(pattern, text, re.IGNORECASE)
        return f"https://{match.group(0)}" if match else None

    def _extract_name(self, lines: List[str]) -> Optional[str]:
        """Extract name (usually first non-empty line)"""
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            if line and not re.search(r'@|http|www|\d{3}', line):
                # Likely a name
                if len(line.split()) <= 4:  # Names usually 2-4 words
                    return line
        return None

    def _identify_sections(self, lines: List[str]) -> Dict[str, List[str]]:
        """Identify and group resume sections"""
        sections = {}
        current_section = "header"
        current_content = []

        for line in lines:
            line_lower = line.lower().strip()

            # Check if this line is a section header
            section_found = None
            for section_name, headers in self.SECTION_HEADERS.items():
                if any(header in line_lower for header in headers):
                    section_found = section_name
                    break

            if section_found:
                # Save previous section
                if current_content:
                    sections[current_section] = current_content
                current_section = section_found
                current_content = []
            else:
                current_content.append(line)

        # Save last section
        if current_content:
            sections[current_section] = current_content

        return sections

    def _parse_summary(self, lines: List[str]) -> str:
        """Parse summary section"""
        return " ".join(line.strip() for line in lines if line.strip())

    def _parse_experiences(self, lines: List[str]) -> List[Dict]:
        """Parse experience section into structured entries"""
        experiences = []
        current_exp = None

        # Patterns for job entries
        date_pattern = r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|Present|\d{4})'

        text = " ".join(lines)
        # Split by date ranges which typically start new entries
        entries = re.split(r'(?=' + date_pattern + r'\s*[-–]\s*' + date_pattern + r')', text)

        for entry in entries:
            if len(entry.strip()) < 20:
                continue

            exp = {
                "title": None,
                "company": None,
                "dates": None,
                "description": ""
            }

            # Try to extract dates
            date_match = re.search(
                date_pattern + r'\s*[-–]\s*' + date_pattern,
                entry
            )
            if date_match:
                exp["dates"] = date_match.group(0)
                entry = entry.replace(exp["dates"], "")

            # First line is usually title/company
            parts = entry.strip().split("\n")
            if parts:
                first_line = parts[0].strip()
                # Try to split title and company
                if " at " in first_line:
                    exp["title"], exp["company"] = first_line.split(" at ", 1)
                elif " - " in first_line:
                    exp["title"], exp["company"] = first_line.split(" - ", 1)
                else:
                    exp["title"] = first_line

                exp["description"] = " ".join(parts[1:]).strip()

            if exp["title"]:
                experiences.append(exp)

        return experiences[:10]  # Limit to 10 experiences

    def _parse_education(self, lines: List[str]) -> List[Dict]:
        """Parse education section"""
        education = []
        text = " ".join(lines)

        # Common degree patterns
        degree_patterns = [
            r"(Bachelor'?s?|Master'?s?|Ph\.?D\.?|MBA|B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?)",
        ]

        # Simple split by institution patterns
        entries = re.split(r'(?=\b(?:University|College|Institute|School)\b)', text, flags=re.IGNORECASE)

        for entry in entries:
            if len(entry.strip()) < 10:
                continue

            edu = {
                "institution": None,
                "degree": None,
                "field": None,
                "year": None
            }

            # Extract year
            year_match = re.search(r'\b(20\d{2}|19\d{2})\b', entry)
            if year_match:
                edu["year"] = int(year_match.group(1))

            # Extract degree
            for pattern in degree_patterns:
                match = re.search(pattern, entry, re.IGNORECASE)
                if match:
                    edu["degree"] = match.group(1)
                    break

            # Institution is usually at the start
            words = entry.strip().split()[:5]
            edu["institution"] = " ".join(words)

            if edu["institution"]:
                education.append(edu)

        return education[:5]

    def _parse_skills(self, lines: List[str]) -> List[str]:
        """Parse skills section into list"""
        skills = []
        text = " ".join(lines)

        # Split by common delimiters
        raw_skills = re.split(r'[,|•·\n]', text)

        for skill in raw_skills:
            skill = skill.strip()
            if skill and len(skill) > 1 and len(skill) < 50:
                # Clean up
                skill = re.sub(r'^[-•·]\s*', '', skill)
                if skill and skill not in skills:
                    skills.append(skill)

        return skills[:50]  # Limit to 50 skills

    def _parse_projects(self, lines: List[str]) -> List[Dict]:
        """Parse projects section"""
        projects = []
        text = " ".join(lines)

        # Split by bullet points or new lines
        entries = re.split(r'[•·]|\n\n', text)

        for entry in entries:
            if len(entry.strip()) < 20:
                continue

            project = {
                "name": None,
                "description": entry.strip()
            }

            # First sentence is usually the name
            parts = entry.strip().split(".")
            if parts:
                project["name"] = parts[0].strip()[:100]
                project["description"] = ".".join(parts[1:]).strip()

            if project["name"]:
                projects.append(project)

        return projects[:10]

    def get_raw_text(self, file_path: str) -> str:
        """Get raw text from a resume file without parsing"""
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self._extract_pdf_text(file_path)
        elif suffix in [".docx", ".doc"]:
            return self._extract_docx_text(file_path)
        elif suffix == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
